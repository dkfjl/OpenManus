import asyncio
from pathlib import Path
from typing import List, Optional

from docx import Document

from app.config import config
from app.exceptions import ToolError
from app.tool.base import BaseTool


_DESCRIPTION = """Create or update a Word .docx document within the workspace.
Provide either a free-form body of text or a list of structured sections (with optional headings and bullet lists).
Set `append` to true to add to an existing document instead of recreating it."""


class WordDocumentTool(BaseTool):
    """Tool for creating structured Word documents."""

    name: str = "word_document"
    description: str = _DESCRIPTION
    parameters: dict = {
        "type": "object",
        "properties": {
            "filepath": {
                "type": "string",
                "description": "Target .docx path (relative to workspace or absolute within it).",
            },
            "document_title": {
                "type": "string",
                "description": "Optional top-level title inserted at the beginning of the document.",
            },
            "body": {
                "type": "string",
                "description": "Optional free-form paragraphs appended before any sections.",
            },
            "sections": {
                "type": "array",
                "description": "Structured sections to add to the document.",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading": {
                            "type": "string",
                            "description": "Heading text for this section.",
                        },
                        "level": {
                            "type": "integer",
                            "minimum": 1,
                            "maximum": 4,
                            "description": "Heading level (1-4). Defaults to 1.",
                        },
                        "content": {
                            "type": "string",
                            "description": "Body text for this section.",
                        },
                        "bullets": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Optional bullet list entries added after the section content.",
                        },
                    },
                    "required": [],
                },
            },
            "append": {
                "type": "boolean",
                "description": "If true and the file exists, add to it instead of recreating it.",
                "default": False,
            },
            "author": {
                "type": "string",
                "description": "Optional author metadata to store in the document properties.",
            },
        },
        "required": ["filepath"],
    }

    async def execute(
        self,
        *,
        filepath: str,
        document_title: Optional[str] = None,
        body: Optional[str] = None,
        sections: Optional[List[dict]] = None,
        append: bool = False,
        author: Optional[str] = None,
        **_: str,
    ):
        if not body and not sections:
            raise ToolError(
                "Provide at least one of `body` or `sections` to write to the document."
            )

        target_path = self._resolve_path(filepath)
        target_path.parent.mkdir(parents=True, exist_ok=True)
        mode = "append" if append and target_path.exists() else "overwrite"

        result = await asyncio.to_thread(
            self._write_document,
            target_path,
            document_title,
            body,
            sections or [],
            append,
            author,
        )

        result["path"] = str(target_path)
        result["mode"] = mode
        return self.success_response(result)

    def _write_document(
        self,
        path: Path,
        document_title: Optional[str],
        body: Optional[str],
        sections: List[dict],
        append: bool,
        author: Optional[str],
    ) -> dict:
        document = Document(str(path)) if append and path.exists() else Document()

        paragraphs_written = 0
        bullets_added = 0

        if document_title:
            document.add_heading(document_title, level=0)

        if body:
            paragraphs_written += self._add_paragraphs(document, body)

        for section in sections:
            heading = (section.get("heading") or "").strip()
            if heading:
                level = self._normalize_level(section.get("level"))
                document.add_heading(heading, level=level)

            content = section.get("content")
            if content:
                paragraphs_written += self._add_paragraphs(document, content)

            for bullet in section.get("bullets") or []:
                text = (bullet or "").strip()
                if not text:
                    continue
                para = document.add_paragraph(text)
                para.style = "List Bullet"
                bullets_added += 1

        if author:
            document.core_properties.author = author

        document.save(str(path))

        return {
            "paragraphs_written": paragraphs_written,
            "sections_written": len(sections),
            "bullets_written": bullets_added,
        }

    @staticmethod
    def _add_paragraphs(document: Document, text: str) -> int:
        normalized = text.replace("\r\n", "\n")
        blocks = [block.strip() for block in normalized.split("\n\n") if block.strip()]
        count = 0
        for block in blocks:
            document.add_paragraph(block)
            count += 1
        return count

    @staticmethod
    def _normalize_level(value: Optional[int]) -> int:
        if value is None:
            return 1
        return max(1, min(4, value))

    @staticmethod
    def _resolve_path(filepath: str) -> Path:
        base = config.workspace_root.resolve()
        candidate = Path(filepath).expanduser()
        if not candidate.is_absolute():
            candidate = base / candidate
        resolved = candidate.resolve()
        if resolved.suffix.lower() != ".docx":
            raise ToolError("Only .docx files are supported by word_document tool.")
        if base not in resolved.parents and resolved != base:
            raise ToolError(
                f"Target path {resolved} is outside of the workspace directory."
            )
        return resolved
