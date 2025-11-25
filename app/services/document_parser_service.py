import asyncio
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

import aiofiles
import pytesseract
from docx import Document
from PIL import Image
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup

from app.config import config
from app.logger import logger


class DocumentParserService:
    """文档解析服务，支持多种文件格式的文本提取"""

    # 支持的文件类型和最大文件大小（10MB）
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.html', '.htm'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    def __init__(self):
        self.upload_dir = config.workspace_root / "uploads"
        self.upload_dir.mkdir(exist_ok=True)

    async def parse_uploaded_files(self, files: List[Any]) -> str:
        """
        解析上传的多个文件，返回合并后的文本内容

        Args:
            files: FastAPI UploadFile 对象列表

        Returns:
            str: 解析后的文本内容，多个文件内容用分隔符连接
        """
        if not files:
            return ""

        if len(files) > 3:
            raise ValueError("最多支持上传3个文件")

        parsed_contents = []

        for file in files:
            try:
                # 验证文件
                await self._validate_file(file)

                # 保存临时文件
                temp_path = await self._save_temp_file(file)

                # 解析文件内容
                content = await self._parse_file(temp_path, file.filename)

                if content.strip():
                    parsed_contents.append(
                        f"=== 从文件 '{file.filename}' 解析的内容 ===\n{content}\n"
                    )

                # 清理临时文件
                await self._cleanup_temp_file(temp_path)

            except Exception as e:
                logger.error(f"解析文件 {file.filename} 失败: {e}")
                # 继续处理其他文件，不因为单个文件失败而中断
                parsed_contents.append(
                    f"=== 文件 '{file.filename}' 解析失败 ===\n错误信息: {str(e)}\n"
                )

        return "\n".join(parsed_contents)

    async def _validate_file(self, file: Any) -> None:
        """验证上传文件"""
        # 检查文件名
        if not file.filename:
            raise ValueError("文件名不能为空")

        # 检查文件扩展名
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS)}")

        # 检查文件大小
        file.file.seek(0, 2)  # 移动到文件末尾
        file_size = file.file.tell()
        file.file.seek(0)  # 重置到文件开头

        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"文件过大: {file_size / (1024*1024):.2f}MB。最大允许: {self.MAX_FILE_SIZE / (1024*1024)}MB")

        if file_size == 0:
            raise ValueError("文件为空")

    async def _save_temp_file(self, file: Any) -> Path:
        """保存上传的文件到临时目录"""
        temp_dir = Path(tempfile.gettempdir()) / "openmanus_uploads"
        temp_dir.mkdir(exist_ok=True)

        temp_path = temp_dir / f"{file.filename}_{os.urandom(8).hex()}"

        async with aiofiles.open(temp_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        logger.info(f"临时文件已保存: {temp_path}")
        return temp_path

    async def _cleanup_temp_file(self, temp_path: Path) -> None:
        """清理临时文件"""
        try:
            if temp_path.exists():
                temp_path.unlink()
                logger.info(f"临时文件已清理: {temp_path}")
        except Exception as e:
            logger.warning(f"清理临时文件失败 {temp_path}: {e}")

    async def _parse_file(self, file_path: Path, filename: str) -> str:
        """根据文件类型解析文件内容"""
        file_ext = Path(filename).suffix.lower()

        try:
            if file_ext == '.pdf':
                return await self._parse_pdf(file_path)
            elif file_ext == '.docx':
                return await self._parse_docx(file_path)
            elif file_ext == '.txt':
                return await self._parse_text(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png']:
                return await self._parse_image(file_path)
            elif file_ext in ['.html', '.htm']:
                return await self._parse_html(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            logger.error(f"解析文件 {filename} 失败: {e}")
            raise

    async def _parse_pdf(self, file_path: Path) -> str:
        """解析PDF文件"""
        def _parse_pdf_sync():
            try:
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    text_content = []

                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(f"第{page_num + 1}页:\n{page_text}")
                        except Exception as e:
                            logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")
                            continue

                    return "\n\n".join(text_content) if text_content else ""
            except Exception as e:
                raise Exception(f"PDF解析失败: {str(e)}")

        return await asyncio.to_thread(_parse_pdf_sync)

    async def _parse_docx(self, file_path: Path) -> str:
        """解析Word文档"""
        def _parse_docx_sync():
            try:
                doc = Document(file_path)
                paragraphs = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)

                # 提取表格内容
                tables_content = []
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    if table_text:
                        tables_content.append("\n".join(table_text))

                content = "\n\n".join(paragraphs)
                if tables_content:
                    content += "\n\n表格内容:\n" + "\n\n".join(tables_content)

                return content
            except Exception as e:
                raise Exception(f"Word文档解析失败: {str(e)}")

        return await asyncio.to_thread(_parse_docx_sync)

    async def _parse_text(self, file_path: Path) -> str:
        """解析纯文本文件"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                return content.strip()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                async with aiofiles.open(file_path, 'r', encoding='gbk') as f:
                    content = await f.read()
                    return content.strip()
            except Exception as e:
                raise Exception(f"文本文件编码解析失败: {str(e)}")
        except Exception as e:
            raise Exception(f"文本文件解析失败: {str(e)}")

    async def _parse_image(self, file_path: Path) -> str:
        """解析图片文件（OCR文字识别）"""
        def _parse_image_sync():
            try:
                image = Image.open(file_path)
                # 使用中文和英文进行OCR
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                return text.strip()
            except Exception as e:
                raise Exception(f"图片OCR识别失败: {str(e)}")

        return await asyncio.to_thread(_parse_image_sync)

    async def _parse_html(self, file_path: Path) -> str:
        """解析HTML文件，提取可读文本"""

        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = await f.read()

        def _extract_text(html: str) -> str:
            soup = BeautifulSoup(html, 'html.parser')
            for tag in soup(['script', 'style', 'noscript']):
                tag.decompose()
            text = soup.get_text(separator='\n')
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return "\n".join(lines)

        return await asyncio.to_thread(_extract_text, html_content)

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """获取支持的文件格式信息"""
        return {
            "supported_extensions": list(self.SUPPORTED_EXTENSIONS),
            "max_file_size_mb": self.MAX_FILE_SIZE / (1024 * 1024),
            "max_file_count": 3,
            "description": {
                ".pdf": "PDF文档",
                ".docx": "Word文档",
                ".txt": "纯文本文件",
                ".jpg/.jpeg": "JPEG图片（OCR识别）",
                ".png": "PNG图片（OCR识别）",
                ".html/.htm": "HTML文档（提取正文）"
            }
        }
